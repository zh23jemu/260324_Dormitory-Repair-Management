from __future__ import annotations

from pathlib import Path
from textwrap import dedent
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# 本脚本用于生成项目文档及配套图例。图例参考用户提供的示例图类别，
# 但内容全部改写为“高校宿舍报修管理系统”的真实模块、流程和数据实体。
ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "project-doc-assets"
MD_PATH = DOCS / "高校宿舍报修管理系统_项目文档.md"
DOCX_PATH = DOCS / "高校宿舍报修管理系统_项目文档.docx"


def ensure_dirs() -> None:
    DOCS.mkdir(exist_ok=True)
    ASSETS.mkdir(exist_ok=True)


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    """加载 Windows 中文字体，保证图例中的中文能够正常显示。"""
    candidates = [
        "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
    ]
    for font_path in candidates:
        if Path(font_path).exists():
            return ImageFont.truetype(font_path, size)
    return ImageFont.load_default()


FONT = load_font(24)
FONT_SM = load_font(18)
FONT_XS = load_font(15)
FONT_BOLD = load_font(28, bold=True)


def text_size(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont) -> tuple[int, int]:
    box = draw.textbbox((0, 0), text, font=font)
    return box[2] - box[0], box[3] - box[1]


def center_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font=FONT_SM, fill="#1f2a44") -> None:
    x1, y1, x2, y2 = box
    lines = text.split("\n")
    heights = [text_size(draw, line, font)[1] for line in lines]
    total_h = sum(heights) + (len(lines) - 1) * 6
    y = y1 + (y2 - y1 - total_h) / 2
    for line, h in zip(lines, heights):
        w, _ = text_size(draw, line, font)
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=font, fill=fill)
        y += h + 6


def box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], text: str, fill="#ffffff", outline="#2f80ed") -> None:
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    center_text(draw, xy, text)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill="#45617d") -> None:
    draw.line([start, end], fill=fill, width=3)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        pts = [(ex, ey), (ex - direction * 14, ey - 8), (ex - direction * 14, ey + 8)]
    else:
        direction = 1 if ey > sy else -1
        pts = [(ex, ey), (ex - 8, ey - direction * 14), (ex + 8, ey - direction * 14)]
    draw.polygon(pts, fill=fill)


def new_canvas(title: str, width: int = 1400, height: int = 850) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (width, height), "#f6f9fc")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((30, 30, width - 30, height - 30), radius=28, fill="#ffffff", outline="#d7e3f3", width=2)
    draw.text((60, 50), title, font=FONT_BOLD, fill="#17324d")
    draw.line((60, 92, width - 60, 92), fill="#d7e3f3", width=2)
    return img, draw


def save_diagram(name: str, img: Image.Image) -> str:
    path = ASSETS / name
    img.save(path)
    return f"project-doc-assets/{name}"


def diagram_flow() -> str:
    img, draw = new_canvas("图 1  宿舍报修业务主流程")
    items = [
        ("学生登录/注册", 80, 160),
        ("选择楼栋、宿舍\n报修类型与期望时间", 330, 160),
        ("提交工单\n上传故障图片", 620, 160),
        ("宿管审核\n通过或驳回", 900, 160),
        ("分配维修员", 1160, 160),
        ("维修员接单", 220, 430),
        ("维修反馈\n耗材上报", 480, 430),
        ("学生评分评价", 750, 430),
        ("工单完成\n统计归档", 1030, 430),
    ]
    coords = []
    for text, x, y in items:
        xy = (x, y, x + 210, y + 100)
        coords.append(xy)
        box(draw, xy, text, fill="#f8fbff")
    for i in range(4):
        arrow(draw, (coords[i][2], 210), (coords[i + 1][0], 210))
    arrow(draw, (1265, 260), (325, 430))
    for i in range(5, 8):
        arrow(draw, (coords[i][2], 480), (coords[i + 1][0], 480))
    draw.text((80, 690), "关键状态：待审核 -> 已驳回/待接单 -> 处理中 -> 待评价 -> 已完成", font=FONT_SM, fill="#45617d")
    return save_diagram("01_repair_flow.png", img)


def diagram_modules() -> str:
    img, draw = new_canvas("图 2  系统功能模块划分")
    center = (585, 170, 815, 250)
    box(draw, center, "高校宿舍报修\n管理系统", fill="#eaf3ff")
    modules = [
        ("公共门户\n公告/工单浏览/统计", 80, 330),
        ("学生端\n报修/工单/评价/留言/论坛", 360, 330),
        ("宿管端\n审核分配/宿舍/学生住宿", 680, 330),
        ("维修员端\n接单/反馈/耗材/统计", 1000, 330),
        ("管理员端\n用户/配置/论坛/耗材/日志", 360, 560),
        ("基础支撑\n认证鉴权/上传/字典/日志", 680, 560),
    ]
    for text, x, y in modules:
        xy = (x, y, x + 250, y + 105)
        box(draw, xy, text, fill="#fbfdff")
        arrow(draw, ((center[0] + center[2]) // 2, center[3]), (x + 125, y))
    return save_diagram("02_module_structure.png", img)


def diagram_er() -> str:
    img, draw = new_canvas("图 3  核心 E-R 实体关系图")
    entities = [
        ("user\n用户/角色/头像/安全问题", 70, 150),
        ("student_profile\n学号/学院/专业/班级/床位", 390, 150),
        ("dorm_building\n楼栋", 730, 150),
        ("dorm_room\n宿舍房间", 1010, 150),
        ("repair_order\n报修工单/地点/状态/维修员", 360, 380),
        ("repair_type\n报修类型", 70, 380),
        ("dorm_facility\n宿舍设施台账", 1010, 380),
        ("repair_flow\n流转记录", 70, 600),
        ("repair_feedback\n维修反馈", 360, 600),
        ("repair_rating\n学生评价", 650, 600),
        ("repair_material\n耗材库存", 920, 600),
        ("material_usage\n耗材使用记录", 1160, 600),
    ]
    boxes = {}
    for text, x, y in entities:
        xy = (x, y, x + 250, y + 105)
        boxes[text.split("\n")[0]] = xy
        box(draw, xy, text, fill="#f9fcff")
    relations = [
        ("user", "student_profile"), ("student_profile", "repair_order"),
        ("dorm_building", "dorm_room"), ("dorm_room", "repair_order"),
        ("dorm_facility", "repair_order"), ("repair_type", "repair_order"),
        ("repair_order", "repair_flow"), ("repair_order", "repair_feedback"),
        ("repair_order", "repair_rating"), ("repair_material", "material_usage"),
        ("repair_order", "material_usage"),
    ]
    for a, b in relations:
        ax = (boxes[a][0] + boxes[a][2]) // 2
        ay = (boxes[a][1] + boxes[a][3]) // 2
        bx = (boxes[b][0] + boxes[b][2]) // 2
        by = (boxes[b][1] + boxes[b][3]) // 2
        arrow(draw, (ax, ay), (bx, by), fill="#8aa0b6")
    return save_diagram("03_er_diagram.png", img)


def diagram_sequence() -> str:
    img, draw = new_canvas("图 4  学生提交报修与审核分配时序图")
    actors = [("学生端 Vue", 120), ("StudentController", 380), ("StudentService", 640), ("SQLite 数据库", 900), ("宿管端", 1160)]
    for name, x in actors:
        box(draw, (x, 150, x + 180, 210), name, fill="#eef6ff")
        draw.line((x + 90, 210, x + 90, 720), fill="#ccd8e6", width=2)
    steps = [
        (0, 1, "提交报修信息/图片/地点", 270),
        (1, 2, "校验字段与登录身份", 340),
        (2, 3, "写入 repair_order/image/flow", 410),
        (3, 2, "返回工单编号", 480),
        (2, 1, "统一响应 ApiResponse", 550),
        (4, 3, "查询待审核工单", 620),
        (4, 3, "分配维修员并更新状态", 690),
    ]
    for src, dst, label, y in steps:
        x1 = actors[src][1] + 90
        x2 = actors[dst][1] + 90
        arrow(draw, (x1, y), (x2, y))
        draw.text((min(x1, x2) + 12, y - 24), label, font=FONT_XS, fill="#30465f")
    return save_diagram("04_sequence_diagram.png", img)


def diagram_architecture() -> str:
    img, draw = new_canvas("图 5  系统部署与技术架构图")
    layers = [
        ("浏览器访问层\nPC 门户首页、统一登录、后台工作台", 90, 150, 1310, 245, "#eaf3ff"),
        ("前端应用层\nVue 3 + Vue Router + Element Plus + Axios，统一端口 5183", 90, 300, 1310, 395, "#f6fbef"),
        ("后端服务层\nSpring Boot 3.3 + Spring Security JWT + Controller/Service/JdbcTemplate，端口 2360", 90, 450, 1310, 545, "#fff8e8"),
        ("数据与文件层\nSQLite 数据库、uploads 文件目录、schema.sql/data.sql 初始化脚本", 90, 600, 1310, 695, "#fdf1f4"),
    ]
    for text, x1, y1, x2, y2, fill in layers:
        box(draw, (x1, y1, x2, y2), text, fill=fill)
    for y in [245, 395, 545]:
        arrow(draw, (700, y), (700, y + 55))
    draw.text((120, 745), "开发/测试/演示环境均可单机部署；生产环境可拆分前端静态服务、后端服务和数据库文件备份目录。", font=FONT_SM, fill="#45617d")
    return save_diagram("05_architecture.png", img)


def diagram_class() -> str:
    img, draw = new_canvas("图 6  后端核心类与分层关系图")
    items = [
        # 类图用于项目文档阅读，不追求列出全部类名；这里使用中文短标签，
        # 避免 Word 缩放图片后英文长类名互相遮挡。
        ("Controller 控制层\n认证/学生/宿管/维修/管理/门户", 80, 160),
        ("Service 业务层\n登录注册/工单/统计/配置", 410, 160),
        ("DTO 入参对象\n登录/报修/派单/反馈", 760, 160),
        ("Common 通用层\n统一响应/业务异常/全局处理", 1070, 160),
        ("Security 安全层\nJWT 令牌/过滤器/角色校验", 260, 430),
        ("JdbcTemplate 数据访问\nSQL 查询/分页/统计/事务", 610, 430),
        ("SQLite 数据表\n用户/工单/宿舍/耗材/日志", 960, 430),
    ]
    coords = []
    for text, x, y in items:
        xy = (x, y, x + 280, y + 115)
        coords.append(xy)
        box(draw, xy, text, fill="#fbfdff")
    arrow(draw, (360, 218), (410, 218))
    arrow(draw, (690, 218), (760, 218))
    arrow(draw, (690, 275), (700, 430))
    arrow(draw, (890, 487), (960, 487))
    arrow(draw, (400, 487), (610, 487))
    arrow(draw, (1210, 275), (1210, 430))
    return save_diagram("06_class_diagram.png", img)


def diagram_use_case() -> str:
    img, draw = new_canvas("图 7  角色用例图")
    actors = [
        ("游客", 90, 170, ["浏览首页", "查看公告", "查看公开工单"]),
        ("学生", 390, 170, ["注册/登录", "提交报修", "查看工单", "评分评价", "留言/论坛"]),
        ("宿管", 720, 170, ["审核工单", "分配维修员", "维护楼栋宿舍", "管理住宿/评价"]),
        ("维修员", 1040, 170, ["接单处理", "反馈完成", "上报耗材", "个人统计"]),
        ("管理员", 550, 520, ["用户管理", "基础配置", "统计分析", "耗材/论坛/日志管理"]),
    ]
    for role, x, y, cases in actors:
        draw.ellipse((x + 80, y, x + 140, y + 60), outline="#2f80ed", width=3)
        draw.line((x + 110, y + 60, x + 110, y + 140), fill="#2f80ed", width=3)
        draw.line((x + 70, y + 90, x + 150, y + 90), fill="#2f80ed", width=3)
        draw.line((x + 110, y + 140, x + 75, y + 200), fill="#2f80ed", width=3)
        draw.line((x + 110, y + 140, x + 145, y + 200), fill="#2f80ed", width=3)
        draw.text((x + 82, y + 215), role, font=FONT_SM, fill="#17324d")
        for i, case in enumerate(cases):
            cy = y + i * 52
            draw.ellipse((x + 190, cy, x + 350, cy + 42), fill="#f8fbff", outline="#8fb7e8", width=2)
            center_text(draw, (x + 190, cy, x + 350, cy + 42), case, font=FONT_XS)
            draw.line((x + 150, y + 100, x + 190, cy + 21), fill="#8aa0b6", width=2)
    return save_diagram("07_use_case.png", img)


def generate_diagrams() -> dict[str, str]:
    return {
        "flow": diagram_flow(),
        "modules": diagram_modules(),
        "er": diagram_er(),
        "sequence": diagram_sequence(),
        "architecture": diagram_architecture(),
        "class": diagram_class(),
        "use_case": diagram_use_case(),
    }


TABLES = [
    ("user", "系统用户表", [
        ("id", "INTEGER", "主键，自增", "用户唯一编号"),
        ("username", "TEXT", "唯一、非空", "登录用户名"),
        ("password", "TEXT", "非空", "登录密码，沿用项目现有明文存储方式"),
        ("real_name", "TEXT", "非空", "真实姓名"),
        ("phone", "TEXT", "可空", "联系电话，前端校验 11 位"),
        ("avatar", "TEXT", "可空", "头像文件路径"),
        ("role", "TEXT", "非空", "student、dorm_admin、repairer、admin"),
        ("work_type_code", "TEXT", "可空", "维修员工种，支持多工种编码"),
        ("password_question", "TEXT", "可空", "非管理员找回密码问题"),
        ("password_answer", "TEXT", "可空", "非管理员找回密码答案"),
        ("status", "TEXT", "默认 enabled", "账号状态"),
        ("created_at/updated_at", "TEXT", "非空", "创建与更新时间"),
    ]),
    ("dorm_building", "宿舍楼栋表", [
        ("id", "INTEGER", "主键", "楼栋编号"),
        ("building_name", "TEXT", "非空", "楼栋名称"),
        ("building_code", "TEXT", "唯一", "楼栋编码"),
        ("gender_type", "TEXT", "可空", "入住性别类型"),
        ("floor_count", "INTEGER", "默认 0", "楼层数"),
        ("remark", "TEXT", "可空", "备注"),
    ]),
    ("dorm_room", "宿舍房间表", [
        ("id", "INTEGER", "主键", "房间编号"),
        ("building_id", "INTEGER", "外键", "所属楼栋"),
        ("room_no", "TEXT", "非空", "宿舍号"),
        ("capacity/current_count", "INTEGER", "默认 0", "容量与当前人数"),
        ("facility_desc", "TEXT", "可空", "房间设施摘要"),
        ("status", "TEXT", "默认 enabled", "房间状态"),
        ("remark", "TEXT", "可空", "备注"),
    ]),
    ("dorm_facility", "宿舍设施台账表", [
        ("id", "INTEGER", "主键", "设施编号"),
        ("room_id", "INTEGER", "外键", "所属宿舍"),
        ("facility_name/facility_type", "TEXT", "非空", "设施名称与类型"),
        ("brand/model_number", "TEXT", "可空", "品牌与型号"),
        ("purchase_date", "TEXT", "可空", "采购日期"),
        ("status", "TEXT", "默认 normal", "设施状态"),
        ("remark", "TEXT", "可空", "备注"),
        ("created_at/updated_at", "TEXT", "非空", "创建与更新时间"),
    ]),
    ("student_profile", "学生档案表", [
        ("id", "INTEGER", "主键", "档案编号"),
        ("user_id", "INTEGER", "唯一外键", "关联用户"),
        ("student_no", "TEXT", "唯一", "学号"),
        ("gender", "TEXT", "可空", "性别"),
        ("college/major/class_name", "TEXT", "可空", "学院、专业、班级"),
        ("building_id/room_id", "INTEGER", "外键", "住宿楼栋与宿舍"),
        ("bed_no", "TEXT", "可空", "床位号"),
        ("created_at/updated_at", "TEXT", "非空", "创建与更新时间"),
    ]),
    ("school_college/school_major/school_class", "学院专业班级表", [
        ("id", "INTEGER", "主键", "基础数据编号"),
        ("college_id/major_id", "INTEGER", "外键", "专业归属学院，班级归属专业"),
        ("college_name/major_name/class_name", "TEXT", "唯一约束", "学院、专业、班级名称"),
        ("sort_no", "INTEGER", "默认 0", "排序号"),
        ("status", "TEXT", "默认 enabled", "启用状态"),
        ("created_at/updated_at", "TEXT", "非空", "创建与更新时间"),
    ]),
    ("repair_type", "报修类型表", [
        ("id", "INTEGER", "主键", "类型编号"),
        ("type_name", "TEXT", "唯一", "报修类型名称"),
        ("sort_no", "INTEGER", "唯一", "排序号，新增时不允许重复"),
        ("status", "TEXT", "默认 enabled", "启用状态"),
    ]),
    ("repair_order", "报修工单主表", [
        ("id/order_no", "INTEGER/TEXT", "主键/唯一", "工单编号"),
        ("student_id", "INTEGER", "外键", "提交学生"),
        ("building_id/room_id/facility_id", "INTEGER", "外键，可空", "提交时选择的维修地点与设施"),
        ("repair_type_id", "INTEGER", "外键", "报修类型"),
        ("title/description", "TEXT", "非空", "标题与故障描述"),
        ("expect_time", "TEXT", "可空", "期望维修时间"),
        ("status", "TEXT", "非空", "工单状态"),
        ("reject_reason", "TEXT", "可空", "驳回原因"),
        ("assigned_repairer_id", "INTEGER", "外键，可空", "分配维修员"),
        ("submitted_at/assigned_at/completed_at", "TEXT", "可空", "提交、派单、完成时间"),
        ("created_at/updated_at", "TEXT", "非空", "创建与更新时间"),
    ]),
    ("repair_order_image", "工单图片表", [
        ("id", "INTEGER", "主键", "图片编号"),
        ("repair_order_id", "INTEGER", "外键", "所属工单"),
        ("image_type", "TEXT", "非空", "报修图片或维修结果图片"),
        ("file_path", "TEXT", "非空", "图片文件路径"),
        ("created_at", "TEXT", "非空", "上传时间"),
    ]),
    ("repair_flow", "工单流转表", [
        ("id", "INTEGER", "主键", "流转编号"),
        ("repair_order_id", "INTEGER", "外键", "所属工单"),
        ("operator_id", "INTEGER", "外键", "操作人"),
        ("from_status/to_status", "TEXT", "可空/非空", "状态变化"),
        ("remark", "TEXT", "可空", "流转说明"),
        ("created_at", "TEXT", "非空", "操作时间"),
    ]),
    ("repair_feedback", "维修反馈表", [
        ("id", "INTEGER", "主键", "反馈编号"),
        ("repair_order_id", "INTEGER", "唯一外键", "所属工单"),
        ("repairer_id", "INTEGER", "外键", "维修员"),
        ("result_desc", "TEXT", "非空", "处理结果说明"),
        ("materials_used", "TEXT", "可空", "耗材描述"),
        ("finish_time", "TEXT", "非空", "完成时间"),
        ("created_at", "TEXT", "非空", "反馈时间"),
    ]),
    ("repair_material/material_usage", "耗材库存与使用表", [
        ("repair_material.id", "INTEGER", "主键", "耗材编号"),
        ("material_name/material_type/unit", "TEXT", "名称唯一", "耗材名称、分类、单位"),
        ("stock_qty/warning_qty", "REAL", "默认 0", "库存数量与预警数量"),
        ("material_usage.repair_order_id", "INTEGER", "外键", "关联工单"),
        ("material_usage.quantity", "REAL", "非空", "本次使用数量，完成反馈时扣减库存"),
        ("created_at/updated_at", "TEXT", "非空", "时间字段"),
    ]),
    ("repair_rating", "维修评价表", [
        ("id", "INTEGER", "主键", "评价编号"),
        ("repair_order_id", "INTEGER", "唯一外键", "一个工单只能评价一次"),
        ("student_id", "INTEGER", "外键", "评价学生"),
        ("score", "INTEGER", "1-5", "服务评分"),
        ("content", "TEXT", "可空", "评价内容"),
        ("created_at", "TEXT", "非空", "评价时间"),
    ]),
    ("announcement", "公告表", [
        ("id", "INTEGER", "主键", "公告编号"),
        ("title/content", "TEXT", "非空", "公告标题与内容"),
        ("image_path", "TEXT", "可空", "公告配图"),
        ("publisher_id", "INTEGER", "外键", "发布者"),
        ("status", "TEXT", "默认 published", "发布状态"),
        ("published_at/created_at", "TEXT", "可空/非空", "发布时间与创建时间"),
    ]),
    ("forum_post/forum_comment", "论坛帖子与评论表", [
        ("forum_post.id", "INTEGER", "主键", "帖子编号"),
        ("student_id", "INTEGER", "外键", "发帖学生"),
        ("title/content/image_path", "TEXT", "内容字段", "帖子标题、正文、配图"),
        ("status", "TEXT", "默认 published", "帖子状态"),
        ("forum_comment.user_id", "INTEGER", "外键", "评论用户"),
        ("forum_comment.content", "TEXT", "非空", "评论内容"),
        ("created_at/updated_at", "TEXT", "非空", "发帖、评论时间"),
    ]),
    ("service_message", "服务留言表", [
        ("id", "INTEGER", "主键", "留言编号"),
        ("student_id", "INTEGER", "外键", "留言学生"),
        ("title/content/image_path", "TEXT", "非空/可空", "留言标题、内容、附图"),
        ("status", "TEXT", "默认 pending", "处理状态"),
        ("reply_content/replied_by/replied_at", "TEXT/INTEGER", "可空", "管理员回复内容、回复人、回复时间"),
        ("created_at/updated_at", "TEXT", "非空", "创建与更新时间"),
    ]),
    ("sys_log/sys_dict", "系统日志与字典表", [
        ("sys_log.user_id", "INTEGER", "外键，可空", "操作用户"),
        ("module_name/operation_type/operation_desc", "TEXT", "非空", "日志模块、类型、描述"),
        ("ip/created_at", "TEXT", "可空/非空", "操作 IP 与时间"),
        ("sys_dict.dict_type/dict_code/dict_name", "TEXT", "非空", "字典类型、编码、名称"),
        ("sort_no/status", "INTEGER/TEXT", "默认值", "排序与启用状态"),
    ]),
]


INTERFACES = [
    ("POST", "/api/auth/login", "统一登录", "username、password", "token、userInfo、role"),
    ("POST", "/api/auth/register", "学生注册", "账号、密码、学号、手机号、学院专业班级、安全问题", "注册结果"),
    ("POST", "/api/auth/forgot-password/question", "获取找回密码问题", "username", "passwordQuestion"),
    ("POST", "/api/auth/forgot-password", "回答问题并重置密码", "username、answer、newPassword", "重置结果"),
    ("GET/PUT", "/api/auth/security-question", "登录后维护安全问题", "question、answer", "保存结果"),
    ("GET", "/api/portal/home", "门户首页数据", "无", "公告、公开工单、服务统计、楼栋排行"),
    ("GET", "/api/portal/repair-orders/{id}", "游客查看公开工单详情", "id", "工单基础信息"),
    ("GET", "/api/portal/announcements", "公开公告分页", "pageNum、pageSize", "公告分页列表"),
    ("GET", "/api/student/forum-posts", "学生论坛列表", "keyword、pageNum、pageSize", "帖子分页列表"),
    ("POST", "/api/student/forum-posts", "发布论坛帖子", "title、content、imagePath", "发布结果"),
    ("POST", "/api/portal/forum-posts/{postId}/comments", "发表评论", "content", "评论结果"),
    ("POST", "/api/student/repair-orders", "提交报修", "楼栋、宿舍、设施、类型、标题、描述、期望时间、图片", "创建结果"),
    ("GET", "/api/student/repair-orders", "我的工单分页", "status、pageNum、pageSize", "工单列表"),
    ("GET", "/api/student/repair-orders/{id}", "学生工单详情", "id", "工单、图片、流程、反馈、评价"),
    ("POST", "/api/student/repair-orders/{id}/rating", "提交评价", "score、content", "评价结果"),
    ("POST", "/api/student/service-messages", "提交服务留言", "title、content、imagePath", "留言结果"),
    ("GET", "/api/dorm-admin/repair-orders", "宿管工单分页筛选", "status、pageNum、pageSize", "工单列表"),
    ("POST", "/api/dorm-admin/repair-orders/query", "工单高级筛选", "关键词、类型、状态、楼栋、维修员、时间范围", "分页结果"),
    ("POST", "/api/dorm-admin/repair-orders/{id}/assign", "审核并分配维修员", "repairerId、remark", "派单结果"),
    ("POST", "/api/dorm-admin/repair-orders/{id}/reject", "驳回工单", "rejectReason", "驳回结果"),
    ("GET/POST/PUT/DELETE", "/api/dorm-admin/buildings", "楼栋管理", "楼栋信息", "增删改查结果"),
    ("GET/POST/PUT/DELETE", "/api/dorm-admin/rooms", "宿舍房间管理", "房间信息", "增删改查结果"),
    ("GET/POST/PUT/DELETE", "/api/dorm-admin/facilities", "设施台账管理", "设施信息", "分页与增删改查结果"),
    ("POST", "/api/repairer/repair-orders/{id}/accept", "维修员接单", "id", "接单结果"),
    ("POST", "/api/repairer/repair-orders/{id}/feedback", "维修完成反馈与耗材上报", "resultDesc、materials、images", "完成结果"),
    ("GET", "/api/repairer/statistics", "维修员个人统计", "dateFrom、dateTo", "完成率、平均处理时长等"),
    ("GET", "/api/admin/statistics/overview", "管理员统计概览", "时间范围", "完成率、满意度、平均评分等"),
    ("GET", "/api/admin/statistics/building-heat", "楼栋工单数量排行", "时间范围", "楼栋排行数据"),
    ("GET/POST/PUT/DELETE", "/api/admin/users", "用户管理", "用户信息、密码、安全问题", "增删改查结果"),
    ("GET/POST/PUT/DELETE", "/api/admin/materials", "耗材管理", "耗材信息", "库存维护结果"),
    ("GET/POST/PUT/DELETE", "/api/admin/school-colleges|school-majors|school-classes", "学院专业班级管理", "基础数据", "增删改查结果"),
]


def table_to_md(headers: Iterable[str], rows: Iterable[Iterable[str]]) -> str:
    headers = list(headers)
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        lines.append("| " + " | ".join(str(cell).replace("\n", "<br>") for cell in row) + " |")
    return "\n".join(lines)


def build_markdown(diagrams: dict[str, str]) -> str:
    db_sections = []
    for table_name, title, fields in TABLES:
        db_sections.append(f"#### {table_name}：{title}\n\n" + table_to_md(["字段名", "类型", "约束", "说明"], fields))
    interface_md = table_to_md(["方法", "接口地址", "用途", "主要入参", "主要出参"], INTERFACES)
    return dedent(f"""
    # 高校宿舍报修管理系统项目文档

    > 文档版本：V1.0  
    > 生成日期：2026-05-09  
    > 适用项目：`C:\\Coding\\260324_Dormitory-Repair-Management`  
    > 说明：本文档中的图例参考附件示例的表达方式，但已按本项目实际业务、技术架构和数据库实体重新绘制。

    ## 一、项目概述

    ### 1. 项目背景与业务现状

    高校宿舍日常运行中，水电、门窗、家具、网络、公共设施等故障具有高频、分散、时效要求高的特点。传统报修方式通常依赖电话、纸质登记、微信群或线下口头反馈，容易出现报修记录不完整、处理进度不透明、维修责任难追踪、满意度难统计、耗材消耗无台账等问题。宿管人员需要在大量信息中进行人工筛选和分派，维修人员缺少统一的接单与反馈入口，管理员也难以及时掌握服务质量、楼栋故障热点和库存耗材变化。

    本项目以“高校宿舍报修管理系统”为建设对象，围绕学生报修、宿管审核派单、维修员处理反馈、学生评价、管理员统计配置等核心流程，构建统一门户和后台工作台。系统支持游客浏览首页公告和公开工单，学生登录后提交报修、发帖留言、评价维修服务；宿管负责工单审核、楼栋宿舍、学生住宿和评价管理；维修员负责接单、反馈、耗材上报和个人统计；管理员负责用户、基础配置、耗材、论坛、留言、公告、日志与统计分析。

    ### 2. 项目目标、核心价值与建设意义

    项目目标是建设一个可演示、可维护、可扩展的宿舍报修业务系统，使报修流程从“人工通知”转为“线上闭环”。系统核心价值包括：提升学生报修便利性，提升宿管派单效率，提升维修人员处理可追踪性，沉淀服务评价和统计数据，支撑楼栋热点分析和耗材库存管理。对毕业设计而言，系统覆盖前后端分离、角色权限、文件上传、分页查询、统计分析、数据库设计和业务流程闭环，能够体现较完整的软件工程实践。

    ### 3. 整体业务流程简要说明

    学生进入统一门户，登录后在报修页面选择楼栋、宿舍、可选设施、报修类型、期望维修时间，填写故障描述并上传图片。工单提交后进入待审核状态，宿管审核通过后分配维修员，也可因信息不完整驳回。维修员接单后进入处理中，完成维修时填写处理结果、上传结果图片并上报耗材使用量，系统同步扣减耗材库存。学生确认服务后进行 1-5 星评分与文字评价，工单最终归档为已完成。系统在首页和后台统计中心展示工单数量、完成率、满意度、楼栋工单数量排行等数据。

    ![宿舍报修业务主流程]({diagrams["flow"]})

    ### 4. 术语定义、缩略词、参考文档

    | 术语 | 说明 |
    | --- | --- |
    | 报修工单 | 学生提交的宿舍故障处理请求，是系统核心业务对象。 |
    | 宿管 | 宿舍管理员，负责审核工单、维护宿舍与学生住宿信息。 |
    | 维修员 | 负责接单、处理工单、反馈维修结果和上报耗材。 |
    | 耗材 | 维修过程中消耗的材料，如灯管、插座、阀门等。 |
    | 统一门户 | 所有角色访问系统的公共首页，提供公告、统计、登录入口和业务入口。 |
    | JWT | JSON Web Token，用于前后端分离场景下的登录态认证。 |

    参考文档包括项目 `README.md`、`PROJECT_SUMMARY.md`、`schema.sql`、`data.sql`、前端路由文件以及各 Controller/Service 源码。

    ## 二、需求分析

    ### 1. 功能性需求

    系统面向游客、学生、宿管、维修员和管理员五类使用者。游客可访问统一首页、浏览公告和公开工单详情。学生可注册登录、找回密码、提交报修、查看我的工单、评价维修服务、维护个人信息和头像、提交服务留言、参与论坛发帖评论、查看维修员信息。宿管可审核和分配工单、维护楼栋宿舍、设施台账、学生住宿和学生班级，管理公告、论坛、留言与评价。维修员可查看可处理工单、接单、完成反馈、上报耗材、查看个人统计、维护头像和多重维修类型。管理员可进行统计分析、用户管理、基础配置、报修类型、学院专业班级、耗材管理、公告论坛留言评价管理和系统日志查询。

    ### 2. 非功能性需求

    系统应满足常规校园宿舍场景的响应速度要求。工单、设施、留言、论坛、用户等列表应分页展示，工单相关查询由后端分页实现，避免数据量增长后影响响应速度。安全方面采用 JWT 登录认证和角色权限控制，限制不同角色访问范围。兼容性方面，前端以 PC 演示为主，同时提供响应式布局以适配手机浏览。可扩展性方面，系统采用 Controller/Service/DTO 分层和统一数据字典，便于扩展新的工单状态、报修类型和维修工种。易用性方面，页面采用统一门户、卡片化布局、中文状态、图片预览、弹窗表单和清晰的错误提示。运维方面，系统采用单体后端、单前端、SQLite 数据库和本地上传目录，便于毕业设计部署与演示。

    ### 3. 约束条件

    技术栈约束为 Vue 3 + Element Plus + Spring Boot + SQLite，保持前后端分离和单体架构，不引入 Redis、消息队列、MinIO、微服务网关等重型组件。运行环境要求 Java 17、Maven、Node.js 和 npm。数据库采用 SQLite，适合单机演示和小规模部署。工期和成本约束要求以最小可交付为前提逐步增强成品感，因此在满足核心流程的同时，优先实现与宿舍报修直接相关的模块。

    ### 4. 业务边界、不包含内容说明

    系统不包含短信或邮箱验证码，不实现线上支付、复杂审批流引擎、物业合同管理、仓库采购审批、匿名论坛、即时聊天、复杂推荐算法和微服务分布式部署。已删除报修知识模块，避免系统主题发散。论坛用于公开交流但浏览、发帖和评论需要登录，评论展示头像、用户名和评论时间。

    ## 三、总体架构设计

    ### 1. 整体架构方案

    系统采用前后端分离的单体架构。前端统一为 `admin-web` 一个 Vue 3 应用，承载公共门户、学生端页面和后台工作台；后端为 `server` Spring Boot 应用，提供 RESTful API、JWT 鉴权、业务服务、文件访问和数据库操作；数据库采用 SQLite，通过 `schema.sql` 和 `data.sql` 初始化表结构与演示数据。

    ![系统部署与技术架构图]({diagrams["architecture"]})

    ### 2. 架构拓扑图、系统模块划分

    系统按访问层、前端应用层、后端服务层、数据与文件层划分。模块上包括公共门户、学生模块、宿管模块、维修员模块、管理员模块和基础支撑模块。

    ![系统功能模块划分]({diagrams["modules"]})

    ### 3. 技术栈选型

    | 层级 | 技术 | 选型说明 |
    | --- | --- | --- |
    | 前端 | Vue 3、Vue Router、Element Plus、Axios | 适合构建 PC 门户和后台管理页面，组件成熟，便于快速开发。 |
    | 后端 | Spring Boot 3.3.12、Spring Security、JdbcTemplate | 单体架构清晰，安全与接口能力完整，SQL 可控。 |
    | 数据库 | SQLite | 部署简单，适合毕业设计演示和单机运行。 |
    | 文件 | 本地 uploads 目录 | 支持头像、公告图、工单图、留言图上传与访问。 |
    | 构建 | Maven、npm/Vite | 分别负责后端和前端构建运行。 |

    ### 4. 部署架构、环境划分

    开发环境在本机启动后端 `2360` 端口和统一前端 `5183` 端口。测试环境可使用同样的单机部署方式，配合独立数据库文件和上传目录。生产环境如需上线，可将前端构建产物部署到 Nginx，后端以 Jar 方式运行，并定期备份 SQLite 数据库和 uploads 目录。

    ## 四、详细模块设计

    ### 1. 各功能模块拆分与职责说明

    | 模块 | 子功能 | 职责说明 |
    | --- | --- | --- |
    | 公共门户 | 首页、公告、公开工单详情、登录入口 | 展示系统能力、服务统计、楼栋工单排行和公告信息。 |
    | 学生模块 | 注册登录、报修、我的工单、评价、留言、论坛、个人中心 | 完成学生从报修到评价的闭环，并支持服务反馈和社区交流。 |
    | 宿管模块 | 工单审核、派单、楼栋宿舍、设施、住宿、班级 | 负责宿舍管理和工单前置审核分配。 |
    | 维修员模块 | 接单、反馈、耗材、统计、个人信息 | 负责工单处理过程和维修结果沉淀。 |
    | 管理员模块 | 统计、用户、配置、耗材、公告、论坛、留言、评价、日志 | 负责系统级配置、数据治理和运营分析。 |
    | 基础支撑 | JWT、统一响应、异常处理、上传、日志 | 为业务模块提供通用能力。 |

    ### 2. 模块间调用关系、交互逻辑

    前端页面通过 Axios 调用后端 REST API，登录成功后保存 JWT Token。路由守卫根据角色控制访问范围。后端 Controller 接收请求后调用 Service，Service 通过 JdbcTemplate 操作 SQLite。工单状态变化时，系统同时写入 `repair_flow` 流转记录；关键操作写入 `sys_log`；维修完成上报耗材时更新 `material_usage` 并扣减 `repair_material.stock_qty`。

    ![后端核心类与分层关系图]({diagrams["class"]})

    ### 3. 核心业务流程详细流程图

    核心流程包括：学生提交报修、宿管审核分配、维修员接单处理、维修反馈与耗材上报、学生评价归档。接口调用时序如下：

    ![学生提交报修与审核分配时序图]({diagrams["sequence"]})

    ### 4. 关键逻辑规则、业务算法说明

    工单状态以中文展示，后端内部状态用于流程判断。未审核工单由宿管审核，驳回后记录驳回原因；审核通过必须分配维修员；维修员接单后才计入个人已接单完成率；完成率计算口径为已接单工单中已完成数量占比。平均处理时长以 `assigned_at` 到 `completed_at` 的小时差计算。耗材上报时要求库存足够，提交成功后扣减库存。评价为 1-5 星，一个工单只能评价一次；管理员或宿管可删除恶意评价，删除不影响工单完成状态，但评价不再展示。报修类型删除时，如果存在未完成工单引用该类型，则禁止删除。手机号注册和修改时前端校验为 11 位。

    ## 五、数据库设计

    ### 1. 数据库选型与存储方案

    项目采用 SQLite 数据库。该方案无需单独安装数据库服务，适合毕业设计和客户本地演示。数据初始化由 Spring Boot 执行 `schema.sql` 和 `data.sql`，上传文件保存在项目根目录 `uploads` 下，数据库中保存文件路径。

    ### 2. E-R 实体关系图

    ![核心 E-R 实体关系图]({diagrams["er"]})

    ### 3. 数据表详细设计

    {"\n\n".join(db_sections)}

    ### 4. 索引设计、分表/缓存策略、数据字典

    系统已为用户角色、工单状态、学生工单、维修员工单、设施工单、工单流转、耗材使用、设施房间、服务留言、论坛帖子状态和论坛评论等字段建立索引。当前系统定位为单机演示与小规模使用，暂不进行分表和外部缓存。若后续工单量显著增长，可优先为 `repair_order.submitted_at`、`repair_order.building_id`、`repair_order.repair_type_id` 增加组合索引，并将图片文件迁移到对象存储。数据字典由 `sys_dict` 管理，包含维修工种、评价指标等配置；报修类型使用独立 `repair_type` 表作为业务主数据。

    ## 六、接口设计

    ### 1. 接口整体规范

    后端接口采用 REST 风格，统一前缀为 `/api`。请求格式主要为 JSON，图片上传使用 multipart/form-data。统一响应结构为 `ApiResponse`，字段包括 `code`、`message` 和 `data`；成功通常返回 `code=200`。需要登录的接口在请求头中携带 `Authorization: Bearer <token>`。接口异常由全局异常处理器转换为统一错误响应。

    ### 2. 前后端/第三方接口清单

    {interface_md}

    ### 3. 核心接口详情

    #### 统一登录 `POST /api/auth/login`

    请求示例：

    ```json
    {{
      "username": "student001",
      "password": "123456"
    }}
    ```

    响应示例：

    ```json
    {{
      "code": 200,
      "message": "success",
      "data": {{
        "token": "jwt-token",
        "userInfo": {{
          "id": 1,
          "username": "student001",
          "role": "student"
        }}
      }}
    }}
    ```

    #### 提交报修 `POST /api/student/repair-orders`

    请求包含楼栋、宿舍、设施、报修类型、标题、描述、期望维修时间和图片路径。后端校验学生身份、报修类型和地点合法性后写入工单主表、图片表和流转表。

    #### 宿管分配维修员 `POST /api/dorm-admin/repair-orders/{{id}}/assign`

    请求包含维修员编号和备注。后端校验工单状态为待审核或待派单，更新 `assigned_repairer_id`、`assigned_at` 和状态，并写入流转记录。

    #### 维修反馈 `POST /api/repairer/repair-orders/{{id}}/feedback`

    请求包含处理结果、维修图片和耗材使用明细。后端校验维修员权限、工单状态和耗材库存，写入 `repair_feedback`、`material_usage`，并扣减 `repair_material.stock_qty`。

    ## 七、UI/交互设计

    ### 1. 页面结构、导航设计

    系统前端采用统一入口。公共层包含 `/home`、`/login`、`/announcements` 等页面。学生功能通过门户导航进入，包括我要报修、我的报修、公告信息、论坛交流、服务留言、维修员信息和个人中心。管理工作台通过左侧菜单组织，管理员、宿管、维修员根据角色看到不同菜单。

    ![角色用例图]({diagrams["use_case"]})

    ### 2. 关键页面原型/效果图

    首页采用宣传门户风格，包含系统 Banner、服务统计、楼栋工单数量排行、公告列表和公开工单列表。报修页面采用 PC 表单卡片布局，按报修地点、故障信息、图片上传、提交确认分组。后台页面采用左侧菜单和内容卡片布局，列表页包含筛选栏、分页、状态标签、详情弹窗和操作按钮。

    ### 3. 交互规则、弹窗、权限控制展示逻辑

    未登录访问受限页面时跳转登录页并携带 redirect 参数。登录页先选择角色，再输入账号密码；选择角色与实际账号角色不一致时拒绝登录。新增和编辑类操作尽量采用弹窗表单，避免页面跳转导致上下文丢失。图片预览支持点击图片或 ESC 关闭，避免遮挡文字。按钮、状态标签和下拉选项均使用中文，提升可读性。

    ## 八、安全设计

    ### 1. 身份认证、权限设计

    系统采用 Spring Security + JWT。后端登录成功后签发 Token，前端保存后在请求头携带。角色包括学生、宿管、维修员和管理员。前端路由守卫控制菜单和页面入口，后端通过 `SecurityUtils.requireRole` 等方式进行接口级权限校验，避免仅依赖前端控制。

    ### 2. 防 SQL 注入、XSS、CSRF、接口防刷

    后端通过 JdbcTemplate 参数绑定执行 SQL，避免字符串拼接造成 SQL 注入。前端对用户输入以文本方式展示，减少 XSS 风险。由于系统采用 JWT 无状态接口，CSRF 防护按前后端分离模式关闭，但受限接口需要 Authorization Token。接口防刷当前采用基础权限限制和表单校验，后续可加入限流拦截器。

    ### 3. 数据加密、敏感信息脱敏、日志审计

    当前项目沿用已有实现，密码和找回答案按普通字段存储，适合毕业设计演示但不建议直接用于生产。后续生产化应改为 BCrypt 哈希密码，并对找回答案加盐哈希。系统已提供 `sys_log` 记录关键操作，包括登录、工单处理、配置维护等，便于审计追踪。

    ## 九、性能与优化设计

    ### 1. 并发、负载、响应时间指标

    在单机演示环境下，目标是普通列表查询 1 秒内响应，工单提交与图片上传在网络正常情况下 3 秒内完成。系统主要面向宿舍维修管理的中低并发场景，不定位为高并发互联网系统。

    ### 2. 缓存方案、异步处理、定时任务设计

    当前版本未引入 Redis 缓存、消息队列和复杂定时任务。统计数据实时查询数据库，保证演示数据变化后立即可见。后续如数据量增长，可对首页统计、公告列表、字典配置等低频变化数据增加本地缓存或 Redis 缓存。

    ### 3. 大数据量/高并发场景优化策略

    工单相关列表采用后端分页，查询参数包括状态、类型、楼栋、维修员、关键词和时间范围，避免一次性返回大量数据。其他表格在数据量较小时可使用 Element Plus 前端分页；若后续规模扩大，应统一迁移到后端分页。数据库层已建立常用索引，后续可根据慢查询增加组合索引。图片文件不直接入库，只保存路径，减少数据库体积。

    ## 十、异常与容错设计

    ### 1. 全局异常捕获、错误处理机制

    后端通过 `GlobalExceptionHandler` 捕获业务异常和系统异常，统一转换为 `ApiResponse`。业务异常通过 `BusinessException` 明确提示，例如无权限、工单状态不允许操作、库存不足、报修类型被未完成工单引用等。前端统一拦截接口错误并弹出中文提示。

    ### 2. 重试机制、降级、熔断、兜底方案

    当前系统为单体本地部署，不引入分布式熔断。对于网络错误，前端提示用户稍后重试；对于图片上传失败，用户可重新上传；对于工单状态冲突，后端以业务异常返回，避免重复提交或越权操作。关键库存扣减应放在同一业务事务中，保证耗材使用和库存变化一致。

    ### 3. 日志收集、报错告警方案

    系统记录业务操作日志到 `sys_log`，后端运行日志由 Spring Boot 输出到控制台。演示环境可通过终端查看错误信息。生产化部署时建议增加文件日志滚动、异常邮件或企业微信告警，并定期备份 SQLite 数据库和上传目录。
    """).strip() + "\n"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[Iterable[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            set_cell_text(cells[i], str(cell_text))
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.color.rgb = RGBColor(23, 50, 77)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10.5)


def add_image(doc: Document, relative_path: str, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(DOCS / relative_path), width=Cm(15.5))
    cp = doc.add_paragraph(caption)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.runs[0].font.name = "微软雅黑"
    cp.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    cp.runs[0].font.size = Pt(9)
    cp.runs[0].font.color.rgb = RGBColor(69, 97, 125)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    styles = doc.styles
    styles["Normal"].font.name = "微软雅黑"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(10.5)


def build_docx(diagrams: dict[str, str]) -> None:
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("高校宿舍报修管理系统项目文档")
    run.bold = True
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(23, 50, 77)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("版本：V1.0    日期：2026-05-09    输出格式：Markdown / Word").font.size = Pt(10)
    doc.add_paragraph()

    add_heading(doc, "一、项目概述", 1)
    add_heading(doc, "1. 项目背景与业务现状", 2)
    add_para(doc, "高校宿舍维修场景具有故障类型多、报修人员分散、处理时效要求高、服务质量需要追踪等特点。传统依赖电话、纸质登记或微信群的方式容易造成信息丢失、处理过程不透明和统计困难。本项目围绕宿舍报修核心业务，建设统一门户和后台工作台，形成学生提交、宿管审核、维修员处理、学生评价、管理员统计的线上闭环。")
    add_heading(doc, "2. 项目目标、核心价值与建设意义", 2)
    add_para(doc, "项目目标是提升报修便利性、派单效率、处理透明度和管理决策能力。系统沉淀工单、设施、耗材、评价、公告、论坛和日志数据，支撑楼栋故障热点分析、满意度统计和耗材库存管理。")
    add_heading(doc, "3. 整体业务流程简要说明", 2)
    add_para(doc, "学生登录后选择维修地点、设施、报修类型和期望时间并提交工单；宿管审核后分配维修员；维修员接单处理并上报耗材；学生评价后工单归档；管理员通过统计中心查看完成率、满意度和楼栋排行。")
    add_image(doc, diagrams["flow"], "图 1  宿舍报修业务主流程")
    add_heading(doc, "4. 术语定义、缩略词、参考文档", 2)
    add_table(doc, ["术语", "说明"], [
        ("报修工单", "学生提交的宿舍故障处理请求，是系统核心业务对象。"),
        ("宿管", "宿舍管理员，负责审核工单和维护住宿数据。"),
        ("维修员", "负责接单、维修反馈和耗材上报。"),
        ("JWT", "前后端分离登录态令牌。"),
    ])

    add_heading(doc, "二、需求分析", 1)
    add_heading(doc, "1. 功能性需求", 2)
    add_para(doc, "系统覆盖游客、学生、宿管、维修员和管理员五类角色。游客可浏览首页公告和公开工单；学生可注册登录、报修、查看工单、评价、留言、论坛交流和维护个人信息；宿管负责审核派单、楼栋宿舍、设施台账、住宿与班级管理；维修员负责接单、反馈、耗材和统计；管理员负责统计、用户、配置、耗材、公告、论坛、留言、评价和日志。")
    add_heading(doc, "2. 非功能性需求", 2)
    add_para(doc, "系统要求列表分页、中文状态展示、图片预览可关闭、PC 端优先美观、移动端可用、接口权限清晰、错误提示明确，并支持后续扩展报修类型、维修工种和评价指标。")
    add_heading(doc, "3. 约束条件", 2)
    add_para(doc, "项目采用 Vue 3、Element Plus、Spring Boot、Spring Security、JdbcTemplate 和 SQLite。系统不引入 Redis、消息队列、复杂工作流或微服务架构，优先满足毕业设计演示和小规模部署。")
    add_heading(doc, "4. 业务边界、不包含内容说明", 2)
    add_para(doc, "系统不包含短信邮箱验证码、线上支付、复杂采购审批、匿名论坛、即时聊天和报修知识模块。论坛浏览和评论均要求登录，评论展示头像、用户名和时间。")

    add_heading(doc, "三、总体架构设计", 1)
    add_heading(doc, "1. 整体架构方案", 2)
    add_para(doc, "系统采用前后端分离单体架构，前端统一为 admin-web Vue 应用，后端为 server Spring Boot 应用，数据库为 SQLite，上传文件保存到本地 uploads 目录。")
    add_image(doc, diagrams["architecture"], "图 2  系统部署与技术架构图")
    add_heading(doc, "2. 架构拓扑图、系统模块划分", 2)
    add_image(doc, diagrams["modules"], "图 3  系统功能模块划分")
    add_heading(doc, "3. 技术栈选型", 2)
    add_table(doc, ["层级", "技术", "说明"], [
        ("前端", "Vue 3、Vue Router、Element Plus、Axios", "统一门户和后台页面。"),
        ("后端", "Spring Boot 3.3.12、Spring Security、JdbcTemplate", "REST 接口、鉴权和 SQL 访问。"),
        ("数据库", "SQLite", "单机部署简单，适合演示。"),
        ("文件", "uploads", "保存头像、公告图、工单图和留言图。"),
    ])
    add_heading(doc, "4. 部署架构、环境划分", 2)
    add_para(doc, "开发环境默认后端端口 2360、前端端口 5183。测试和演示环境可复用单机部署方式；生产环境可拆分前端静态服务、后端 Jar 服务、数据库文件和上传目录备份。")

    add_heading(doc, "四、详细模块设计", 1)
    add_heading(doc, "1. 各功能模块拆分与职责说明", 2)
    add_table(doc, ["模块", "职责"], [
        ("公共门户", "展示系统能力、服务统计、楼栋工单排行、公告和公开工单。"),
        ("学生模块", "报修、工单、评价、留言、论坛、个人中心和安全问题。"),
        ("宿管模块", "审核派单、楼栋宿舍、设施台账、学生住宿和评价管理。"),
        ("维修员模块", "接单、完成反馈、耗材上报、个人统计和多工种维护。"),
        ("管理员模块", "统计、用户、配置、耗材、公告、论坛、留言、评价和日志。"),
    ])
    add_heading(doc, "2. 模块间调用关系、交互逻辑", 2)
    add_para(doc, "前端通过 Axios 调用后端接口，后端 Controller 调用 Service，Service 通过 JdbcTemplate 操作 SQLite。工单状态变化写入 repair_flow，关键操作写入 sys_log。")
    add_image(doc, diagrams["class"], "图 4  后端核心类与分层关系图")
    add_heading(doc, "3. 核心业务流程详细流程图", 2)
    add_image(doc, diagrams["sequence"], "图 5  学生提交报修与审核分配时序图")
    add_heading(doc, "4. 关键逻辑规则、业务算法说明", 2)
    add_para(doc, "工单从待审核流转到待接单、处理中、待评价和已完成。维修员完成率只统计已接单工单，平均处理时长按 assigned_at 到 completed_at 计算。耗材上报成功后扣减库存；报修类型若被未完成工单引用则禁止删除；评价删除不影响工单完成状态。")

    add_heading(doc, "五、数据库设计", 1)
    add_heading(doc, "1. 数据库选型与存储方案", 2)
    add_para(doc, "数据库采用 SQLite，初始化脚本位于 server/src/main/resources/schema.sql 和 data.sql。图片和附件保存在 uploads 目录，数据库保存相对路径。")
    add_heading(doc, "2. E-R 实体关系图", 2)
    add_image(doc, diagrams["er"], "图 6  核心 E-R 实体关系图")
    add_heading(doc, "3. 数据表详细设计", 2)
    for table_name, title, fields in TABLES:
        add_heading(doc, f"{table_name}：{title}", 3)
        add_table(doc, ["字段名", "类型", "约束", "说明"], fields)
    add_heading(doc, "4. 索引设计、分表/缓存策略、数据字典", 2)
    add_para(doc, "系统已对用户角色、工单状态、学生、维修员、设施、流转、耗材、论坛和留言等常用查询字段建立索引。当前不做分表和外部缓存，后续可根据慢查询增加组合索引。sys_dict 管理维修工种和评价指标，repair_type 作为报修类型业务主数据。")

    add_heading(doc, "六、接口设计", 1)
    add_heading(doc, "1. 接口整体规范", 2)
    add_para(doc, "接口统一以 /api 为前缀，JSON 作为主要请求和响应格式，文件上传采用 multipart/form-data。统一响应结构为 code、message、data；需要登录的接口通过 Authorization Bearer Token 鉴权。")
    add_heading(doc, "2. 前后端接口清单", 2)
    add_table(doc, ["方法", "接口地址", "用途", "主要入参", "主要出参"], INTERFACES)
    add_heading(doc, "3. 核心接口详情", 2)
    add_para(doc, "统一登录接口 POST /api/auth/login 返回 JWT 和用户角色；提交报修接口 POST /api/student/repair-orders 写入工单、图片和流转记录；宿管分配接口 POST /api/dorm-admin/repair-orders/{id}/assign 更新维修员和状态；维修反馈接口 POST /api/repairer/repair-orders/{id}/feedback 写入维修结果、图片和耗材使用，并同步扣减库存。")

    add_heading(doc, "七、UI/交互设计", 1)
    add_heading(doc, "1. 页面结构、导航设计", 2)
    add_para(doc, "统一门户作为系统入口，学生和管理员均从同一登录页进入。学生侧以 PC 表单和卡片为主，后台侧采用左侧菜单和顶部栏。")
    add_image(doc, diagrams["use_case"], "图 7  角色用例图")
    add_heading(doc, "2. 关键页面原型/效果图", 2)
    add_para(doc, "首页包含 Banner、服务统计、楼栋工单排行、公告和公开工单。报修页按地点、故障信息、图片上传和提交确认分组。后台列表页包含筛选栏、分页、状态标签、详情弹窗和操作按钮。")
    add_heading(doc, "3. 交互规则、弹窗、权限控制展示逻辑", 2)
    add_para(doc, "未登录访问受限页跳转登录；登录页先选角色再登录；新增编辑采用弹窗；图片预览支持点击图片或 ESC 关闭；所有状态和按钮尽量使用中文。")

    add_heading(doc, "八、安全设计", 1)
    add_heading(doc, "1. 身份认证、权限设计", 2)
    add_para(doc, "系统使用 Spring Security 与 JWT 鉴权。前端根据路由 meta.roles 控制页面访问，后端继续做角色校验，避免越权。")
    add_heading(doc, "2. 防 SQL 注入、XSS、CSRF、接口防刷", 2)
    add_para(doc, "数据库访问使用 JdbcTemplate 参数绑定，降低 SQL 注入风险。前端文本展示减少 XSS 风险。JWT 无状态接口关闭传统 CSRF，但所有受限接口均需 Token。")
    add_heading(doc, "3. 数据加密、敏感信息脱敏、日志审计", 2)
    add_para(doc, "当前密码和找回答案沿用演示项目普通字段存储方式；生产化时建议改为哈希存储。sys_log 记录关键业务操作，支持审计追踪。")

    add_heading(doc, "九、性能与优化设计", 1)
    add_heading(doc, "1. 并发、负载、响应时间指标", 2)
    add_para(doc, "系统面向校园宿舍中低并发场景，普通查询目标 1 秒内响应，图片上传和工单提交目标 3 秒内完成。")
    add_heading(doc, "2. 缓存方案、异步处理、定时任务设计", 2)
    add_para(doc, "当前不引入 Redis 或消息队列，统计实时查询数据库，保证演示数据即时变化。后续可对首页统计和字典配置增加缓存。")
    add_heading(doc, "3. 大数据量/高并发场景优化策略", 2)
    add_para(doc, "工单相关列表采用后端分页，其他小规模列表可使用 Element Plus 前端分页。后续可增加组合索引、对象存储和静态资源缓存。")

    add_heading(doc, "十、异常与容错设计", 1)
    add_heading(doc, "1. 全局异常捕获、错误处理机制", 2)
    add_para(doc, "GlobalExceptionHandler 将业务异常和系统异常转换为统一响应，前端弹出中文错误提示。")
    add_heading(doc, "2. 重试机制、降级、熔断、兜底方案", 2)
    add_para(doc, "系统为单体部署，不引入熔断。网络失败提示用户重试，状态冲突由后端业务异常兜底，库存扣减与维修反馈保持同一业务操作。")
    add_heading(doc, "3. 日志收集、报错告警方案", 2)
    add_para(doc, "业务日志写入 sys_log，运行日志输出到控制台。生产化建议增加日志文件滚动、异常告警和数据库/上传目录定期备份。")

    doc.save(DOCX_PATH)


def main() -> None:
    ensure_dirs()
    diagrams = generate_diagrams()
    MD_PATH.write_text(build_markdown(diagrams), encoding="utf-8")
    build_docx(diagrams)
    print(f"Markdown: {MD_PATH}")
    print(f"Word: {DOCX_PATH}")
    print(f"Assets: {ASSETS}")


if __name__ == "__main__":
    main()
