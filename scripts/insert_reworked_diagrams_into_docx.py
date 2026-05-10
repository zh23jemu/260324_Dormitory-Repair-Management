from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph


# 将 docs/project-doc-assets/reworked-diagrams 目录下已重绘的图片，
# 插入到现有 Word 文档的对应章节。脚本只修改目标 docx，不生成新图片。
ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "docs" / "高校宿舍报修管理系统_项目文档.docx"
IMG_DIR = ROOT / "docs" / "project-doc-assets" / "reworked-diagrams"

IMAGES = {
    "architecture": IMG_DIR / "图6_平台系统架构图.png",
    "uml": IMG_DIR / "图7_UML类图.png",
    "sequence": IMG_DIR / "图5_时序图.png",
    "er": IMG_DIR / "图4_系统E-R图.png",
    "use_case": IMG_DIR / "系统用例图.png",
}


def paragraph_has_picture(paragraph) -> bool:
    return bool(paragraph._element.xpath(".//w:drawing"))


def clear_paragraph(paragraph) -> None:
    """清空段落中的全部 run，用于复用文档里原本预留的空白图片段落。"""
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def add_picture_to_paragraph(paragraph, image_path: Path, width_cm: float) -> None:
    clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))


def style_caption(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.font.name = "微软雅黑"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(69, 97, 125)


def set_caption(paragraph, text: str) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.font.name = "微软雅黑"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(69, 97, 125)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def insert_paragraph_after(paragraph, text: str = ""):
    """在指定段落后插入新段落，返回新段落对象。"""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def find_paragraph(doc: Document, exact_text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == exact_text:
            return paragraph
    raise RuntimeError(f"未找到段落：{exact_text}")


def find_previous_picture_paragraph(doc: Document, caption_text: str):
    caption_index = None
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == caption_text:
            caption_index = index
            break
    if caption_index is None:
        raise RuntimeError(f"未找到图题：{caption_text}")
    for index in range(caption_index - 1, -1, -1):
        paragraph = doc.paragraphs[index]
        if paragraph_has_picture(paragraph) or not paragraph.text.strip():
            return paragraph, doc.paragraphs[caption_index]
        if paragraph.text.strip().startswith(("一、", "二、", "三、", "四、", "五、", "六、", "七、", "八、", "九、", "十、")):
            break
    raise RuntimeError(f"未找到图题前的图片/空白段落：{caption_text}")


def main() -> None:
    for image in IMAGES.values():
        if not image.exists():
            raise FileNotFoundError(image)

    doc = Document(DOCX)

    # 1. 平台系统架构图：文档原来只有架构说明，没有单独配图。
    arch_context = find_paragraph(
        doc,
        "系统采用前后端分离单体架构，前端统一为 admin-web Vue 应用，后端为 server Spring Boot 应用，数据库为 SQLite，上传文件保存到本地 uploads 目录。"
    )
    arch_context.text = (
        "系统采用前后端分离单体架构，前端统一为 admin-web Vue 应用，后端为 server Spring Boot 应用，"
        "数据库为 SQLite，上传文件保存到本地 uploads 目录。图 2 从客户端、角色对象、前端应用层、"
        "后端服务层、数据表层、数据库文件与运行环境等维度展示系统整体架构。"
    )
    arch_caption = insert_paragraph_after(arch_context, "图 2  平台系统架构图")
    arch_pic = insert_paragraph_after(arch_context)
    add_picture_to_paragraph(arch_pic, IMAGES["architecture"], 15.2)
    set_caption(arch_caption, "图 2  平台系统架构图")
    # 插入顺序修正：addnext 连续插入会后插在前，需要确保图片在图题前。
    arch_context._p.addnext(arch_pic._p)
    arch_pic._p.addnext(arch_caption._p)

    # 2. UML 类图：插入到“模块间调用关系、交互逻辑”后。
    uml_pic, uml_caption = find_previous_picture_paragraph(doc, "图 4  后端核心类与分层关系图")
    add_picture_to_paragraph(uml_pic, IMAGES["uml"], 14.6)
    set_caption(uml_caption, "图 4  后端核心 UML 类图")

    # 3. 时序图：插入到“核心业务流程详细流程图”位置。
    seq_pic, seq_caption = find_previous_picture_paragraph(doc, "图 5  学生提交报修与审核分配时序图")
    add_picture_to_paragraph(seq_pic, IMAGES["sequence"], 15.2)
    set_caption(seq_caption, "图 5  学生提交报修与审核分配时序图")

    # 4. E-R 图：插入到数据库设计章节。
    er_pic, er_caption = find_previous_picture_paragraph(doc, "图 6  核心 E-R 实体关系图")
    add_picture_to_paragraph(er_pic, IMAGES["er"], 15.2)
    set_caption(er_caption, "图 6  核心 E-R 实体关系图")

    # 5. 系统用例图：插入到 UI/交互设计章节，并更新图题。
    use_pic, use_caption = find_previous_picture_paragraph(doc, "图 7  角色用例图")
    add_picture_to_paragraph(use_pic, IMAGES["use_case"], 15.2)
    set_caption(use_caption, "图 7  系统用例图")

    # 统一整理已有图题样式。
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("图 "):
            style_caption(paragraph)

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
